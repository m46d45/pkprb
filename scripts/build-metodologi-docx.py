#!/usr/bin/env python3
"""Generate shareable methodology DOCX from the PKPRB scoring spec."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(0x1F, 0x1A, 0x16)
MUTED = RGBColor(0x6B, 0x63, 0x58)
TEAL = RGBColor(0x3D, 0x7F, 0x92)
PAPER = "F3EEE4"
LINE = "D9D1C5"
HEADER_BG = "1F1A16"
ROW_ALT = "F7F3EC"
FORMULA_BG = "EDE7DC"

DISCIPLINES = [
    ("sipil", "Teknik Sipil"),
    ("arsitektur", "Arsitektur"),
    ("pwk", "Planologi"),
    ("geologi", "Teknik Geologi"),
    ("lingkungan", "Teknik Lingkungan"),
    ("kelautan", "Teknik Kelautan"),
    ("bencana", "Kebencanaan"),
    ("multidisiplin", "Multidisiplin"),
]

WEIGHTS = {
    "Gempabumi": {
        "sipil": 1.00, "arsitektur": 0.55, "pwk": 0.40, "geologi": 0.90,
        "lingkungan": 0.25, "kelautan": 0.25, "bencana": 0.85, "multidisiplin": 0.70,
    },
    "Tsunami": {
        "sipil": 0.70, "arsitektur": 0.35, "pwk": 0.80, "geologi": 0.65,
        "lingkungan": 0.40, "kelautan": 0.90, "bencana": 0.90, "multidisiplin": 0.75,
    },
    "Banjir": {
        "sipil": 0.85, "arsitektur": 0.30, "pwk": 0.85, "geologi": 0.35,
        "lingkungan": 0.75, "kelautan": 0.35, "bencana": 0.85, "multidisiplin": 0.80,
    },
    "Tanah longsor": {
        "sipil": 0.90, "arsitektur": 0.25, "pwk": 0.55, "geologi": 0.95,
        "lingkungan": 0.40, "kelautan": 0.15, "bencana": 0.80, "multidisiplin": 0.70,
    },
    "Likuefaksi": {
        "sipil": 1.00, "arsitektur": 0.25, "pwk": 0.35, "geologi": 0.90,
        "lingkungan": 0.20, "kelautan": 0.15, "bencana": 0.70, "multidisiplin": 0.55,
    },
    "Gunung api": {
        "sipil": 0.55, "arsitektur": 0.30, "pwk": 0.50, "geologi": 0.95,
        "lingkungan": 0.40, "kelautan": 0.10, "bencana": 0.85, "multidisiplin": 0.70,
    },
    "Karhutla": {
        "sipil": 0.25, "arsitektur": 0.15, "pwk": 0.45, "geologi": 0.25,
        "lingkungan": 0.70, "kelautan": 0.05, "bencana": 0.80, "multidisiplin": 0.75,
    },
    "Komposit infrastruktur": {
        "sipil": 0.85, "arsitektur": 0.40, "pwk": 0.55, "geologi": 0.70,
        "lingkungan": 0.45, "kelautan": 0.40, "bencana": 0.85, "multidisiplin": 0.80,
    },
}

CENTERS = [
    ("Pusat Studi Bencana (PSB)", "Universitas Andalas", "Sumatera Barat",
     "Megathrust Sumatera, gempa-tsunami Padang; berdiri 2007."),
    ("Tsunami and Disaster Mitigation Research Center (TDMRC)", "Universitas Syiah Kuala", "Aceh",
     "Tsunami dan gempa pasca-2004; laboratorium flume tsunami; jejaring internasional."),
    ("Pusat Studi Bencana (PSBA)", "Universitas Gadjah Mada", "Daerah Istimewa Yogyakarta",
     "Multidisiplin pasca gempa Yogya 2006; pendidikan dan pengabdian."),
    ("Pusat Studi Manajemen Bencana (PSMB)", "Universitas Pembangunan Nasional Veteran Yogyakarta",
     "Daerah Istimewa Yogyakarta", "Manajemen risiko dan geobencana."),
    ("Pusat Studi Bencana Unhas", "Universitas Hasanuddin", "Sulawesi Selatan",
     "Banjir, longsor, dan pesisir/DAS Sulawesi Selatan."),
    ("PUI Gambut & Kebencanaan (PUI-GK)", "Universitas Riau", "Riau",
     "Karhutla dan tata kelola gambut."),
    ("Disaster Risk Reduction Center (DRRC UI)", "Universitas Indonesia", "Jawa Barat",
     "PRB, krisis, kesehatan lingkungan."),
    ("P3B LPPM UNS", "Universitas Sebelas Maret", "Jawa Tengah",
     "Kebijakan dan pendidikan kebencanaan."),
    ("Puslit MKPI ITS", "Institut Teknologi Sepuluh Nopember", "Jawa Timur",
     "Mitigasi bencana dan perubahan iklim."),
    ("Research Center for Disaster Mitigation (RCDM)", "Institut Teknologi Bandung", "Jawa Barat",
     "Seismologi, vulkanologi, kontribusi PuSGeN; dampak nasional."),
    ("Disaster Risk Reduction Center (DiRReC)", "Universitas Islam Indonesia",
     "Daerah Istimewa Yogyakarta", "PRB dan pengabdian masyarakat."),
    ("PSMPB UAD", "Universitas Ahmad Dahlan", "Daerah Istimewa Yogyakarta",
     "Mitigasi, edukasi, ketangguhan."),
    ("Pusat Studi Nalodo Internasional", "Universitas Tadulako", "Sulawesi Tengah",
     "Likuefaksi/Nalodo pasca Palu 2018."),
]


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="D9D1C5"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("PKPRB  ·  pkprb.vercel.app/metodologi  ·  P2MI FTSL ITB 2026")
    set_run_font(run, size=8, color=MUTED)
    p.add_run("    ")
    # page number field
    run2 = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run2._r.append(fld)
    run3 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run3._r.append(instr)
    run4 = p.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run4._r.append(fld2)
    for r in (run2, run3, run4):
        set_run_font(r, size=8, color=MUTED)


def p_style(paragraph, space_after=8, space_before=0):
    pf = paragraph.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK
        run.font.name = "Cambria"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    p_style(h, space_after=8, space_before=16 if level == 1 else 12)
    return h


def body(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p_style(p)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def mixed(doc, parts, size=11):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p_style(p)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p_style(p, space_after=4)
        if isinstance(item, str):
            run = p.add_run(item)
            set_run_font(run)
        else:
            for text, bold, italic in item:
                run = p.add_run(text)
                set_run_font(run, bold=bold, italic=italic)


def formula(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, FORMULA_BG)
    set_cell_border(cell, "C9C0B3")
    p = cell.paragraphs[0]
    p_style(p, space_after=0)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=10)
    doc.add_paragraph()


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = True
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, HEADER_BG)
        set_cell_border(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p_style(p, space_after=0)
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xF3, 0xEE, 0xE4))
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            if r_i % 2 == 1:
                shade(cell, ROW_ALT)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p_style(p, space_after=0)
            run = p.add_run(str(val))
            set_run_font(run, size=9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def callout(doc, title, body_text):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell, "E8D9D0")
    set_cell_border(cell, "C45C48")
    p = cell.paragraphs[0]
    p_style(p, space_after=4)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True)
    p2 = cell.add_paragraph()
    p_style(p2, space_after=0)
    run2 = p2.add_run(body_text)
    set_run_font(run2, size=10)
    doc.add_paragraph()


def dec(n):
    return f"{n:.2f}".replace(".", ",")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    add_footer(section)

    kicker = doc.add_paragraph()
    p_style(kicker, space_after=2)
    run = kicker.add_run("PROTOTIPE  ·  20 AGUSTUS 2026  ·  P2MI MULTIDISIPLIN FTSL ITB")
    set_run_font(run, size=9, bold=True, color=TEAL)

    title = doc.add_paragraph()
    p_style(title, space_after=4, space_before=0)
    run = title.add_run("Metodologi PKPRB")
    set_run_font(run, name="Cambria", size=26, bold=True)

    sub = doc.add_paragraph()
    p_style(sub, space_after=12)
    run = sub.add_run("Peta Keselarasan Pendidikan dan Risiko Bencana")
    set_run_font(run, name="Cambria", size=14, italic=True, color=MUTED)

    body(
        doc,
        "Dokumen ini menjelaskan rumus yang benar-benar dijalankan peta, "
        "supaya tim dapat meninjau asumsi, bobot, dan celah data. Versi web "
        "yang sama: https://pkprb.vercel.app/metodologi — peta: https://pkprb.vercel.app",
    )

    callout(
        doc,
        "Bukan peta kesiapan atau ketangguhan daerah.",
        "PKPRB memetakan keselarasan antara profil risiko dan dukungan pendidikan "
        "tinggi (prodi, penelitian, pengabdian). Tidak mengukur BPBD, SNI, stok "
        "insinyur, atau mutu lulusan di lapangan.",
    )

    heading(doc, "Masukan yang kami butuhkan")
    body(
        doc,
        "Silakan uji peta lalu komentari butir berikut. Geser sliders: peta dan rumus ini sama.",
    )
    bullets(
        doc,
        [
            "Matriks bobot disiplin × bahaya (prior Delphi). Apakah sipil harus 1,00 di gempa? Apakah Planologi 0,40 terlalu rendah?",
            "Jenjang S1/S2/S3 default semua 1,00. Perlukah S2/S3 lebih tinggi?",
            "IABEE dipetakan ke akreditasi Internasional, bukan bonus di atas Unggul. Setuju?",
            "Hanya ITB, UI, UGM, ITS yang disebut PT internasional untuk spillover. Tambah IPB, Unpad, Unhas, atau lainnya?",
            "Besaran spillover 12 / 8 / 4% sepulau dan 6 / 3% antar-pulau. Terlalu besar, terlalu kecil, atau arahnya salah?",
            "Skala warna mutlak: risiko 0–100 (komposit 0–200), pendidikan 0–15 per juta. Cap 15 terlalu ketat?",
            "Inventaris Planologi, geologi, arsitektur, lingkungan, kelautan, multidisiplin masih kurasi awal — bukan direktori BAN-PT lengkap.",
            "Label kuadran: Senjang, Selaras, Berlebih, Relevan.",
        ],
        numbered=True,
    )

    heading(doc, "Cara membaca peta")
    bullets(
        doc,
        [
            [
                ("Pendidikan", True, False),
                (" — indeks per juta penduduk (IDPKI), skala warna mutlak 0–15.", False, False),
            ],
            [
                ("Risiko", True, False),
                (" — komposit bergaya IRBI atau satu jenis bahaya. Skala mutlak 0–100 (0–200 untuk komposit).", False, False),
            ],
            [
                ("Keselarasan", True, False),
                (" — matriks 3×3 tertil risiko × tertil pendidikan di antara 38 provinsi. Ini satu-satunya tampilan yang memang relatif antarprovinsi.", False, False),
            ],
        ],
        numbered=True,
    )

    heading(doc, "1. Skor prodi")
    body(doc, "Tidak ada toggle. Bobot 0 = tidak dihitung. Setiap program studi:")
    formula(
        doc,
        "E_prodi = w(disiplin, bahaya)  ×  w(jenjang)  ×  w(akreditasi_prodi)",
    )

    heading(doc, "Jenjang", level=2)
    body(
        doc,
        "Default sementara sama, agar perbedaan antar-disiplin mudah dibaca. D4 memakai bobot S1.",
    )
    table(doc, ["Jenjang", "Default"], [["S1 (dan D4)", "1,00"], ["S2", "1,00"], ["S3", "1,00"]])

    heading(doc, "Akreditasi prodi", level=2)
    body(
        doc,
        "IABEE General atau Provisional → Internasional. Tidak ditumpuk di atas Unggul.",
    )
    table(
        doc,
        ["Peringkat prodi", "Default"],
        [
            ["Internasional (IABEE)", "1,00"],
            ["Unggul atau A", "0,90"],
            ["Baik Sekali atau B", "0,80"],
            ["Baik, C, atau terakreditasi lain", "0,70"],
        ],
    )

    heading(doc, "2. Bobot disiplin × bahaya")
    body(
        doc,
        "Prior 0–1, berganti otomatis saat jenis bahaya dipilih, lalu dapat digeser. "
        "Sipil tinggi pada gempa dan likuefaksi; Planologi pada tsunami dan banjir; "
        "geologi pada longsor dan gunung api; kelautan pada tsunami.",
    )
    w_headers = ["Bahaya"] + [label for _, label in DISCIPLINES]
    w_rows = []
    for hazard, vals in WEIGHTS.items():
        w_rows.append([hazard] + [dec(vals[k]) for k, _ in DISCIPLINES])
    table(doc, w_headers, w_rows)

    heading(doc, "3. Penelitian: pusat studi")
    body(doc, "Slider w_pusat (default 1). Nol = pusat studi tidak dihitung.")
    formula(doc, "R_pusat = w_pusat × 1,35 × kematangan × kesesuaian × nasional")
    table(
        doc,
        ["Faktor", "Nilai"],
        [
            ["Kematangan: anchor (TDMRC, PSBA, RCDM, PSB Unand, Nalodo)", "1,35"],
            ["Kematangan: PUI", "1,20"],
            ["Kematangan: standar", "1,00"],
            ["Bahaya cocok, atau mode komposit", "1,00"],
            ["Bahaya tidak tercantum di fokus pusat", "0,20"],
            ["Pusat berjejaring nasional", "1,20"],
            ["Pusat lokal", "1,00"],
        ],
    )

    heading(doc, "4. Pengabdian: layanan kepakaran")
    body(
        doc,
        "Dipisah dari penelitian. Hanya pusat yang punya rekam PkM. Slider w_kepakaran "
        "default 1. Basis data masih proxy dari pusat studi, bukan direktori layanan "
        "kepakaran tersendiri.",
    )
    formula(doc, "K_kepakaran = w_kepakaran × 1,10 × kematangan × kesesuaian")
    table(
        doc,
        ["Kematangan", "Pengali"],
        [["Anchor", "1,20"], ["PUI", "1,10"], ["Standar", "1,00"]],
    )

    heading(doc, "5. Spillover antarprovinsi")
    mixed(
        doc,
        [
            ("Mengikuti ", False, False),
            ("akreditasi perguruan tinggi", True, False),
            (", bukan peringkat prodi. Provinsi asal tetap 100% skornya; tetangganya "
             "mendapat tambahan. Slider w_spill (default 1) menaik-turunkan semua persentase.", False, False),
        ],
    )
    table(
        doc,
        ["Akreditasi PT", "Pulau yang sama", "Pulau lain"],
        [
            ["Internasional (ITB, UI, UGM, ITS)", "12% × w_spill", "6% × w_spill"],
            ["Unggul", "8% × w_spill", "3% × w_spill"],
            ["Baik Sekali", "4% × w_spill", "0"],
            ["Baik", "0", "0"],
        ],
    )

    heading(doc, "6. Kapasitas dan IDPKI")
    formula(
        doc,
        "Kapasitas = Σ E_prodi + Σ R_pusat + Σ K_kepakaran\n"
        "            (masing-masing termasuk spillover)\n\n"
        "IDPKI = Kapasitas / (jumlah penduduk / 1.000.000)",
    )
    body(
        doc,
        "Per juta penduduk, supaya Jawa tidak otomatis “tinggi” hanya karena jumlah "
        "kampus. Tab Pendidikan mewarnai skala mutlak 0–15. Nilai di atas 15 tetap "
        "dihitung, warnanya menempel di ujung teal.",
    )

    heading(doc, "7. Risiko")
    body(
        doc,
        "Komposit dikalibrasi ke angka publik IRBI 2024: Maluku 161,5; Maluku Utara "
        "145,09; DKI Jakarta 59,29. Skor per bahaya (gempabumi, tsunami, banjir, "
        "longsor, likuefaksi, gunung api, karhutla) adalah profil relatif prototipe "
        "— bukan sel resmi IRBI/BNPB.",
    )
    body(
        doc,
        "Warna Risiko: 0–100 untuk bahaya tunggal, 0–200 untuk komposit. Tidak "
        "distretch ke provinsi terendah–tertinggi.",
    )

    heading(doc, "8. Tertil dan matriks 3×3")
    body(
        doc,
        "Kelas 0 / 1 / 2 dihitung ulang setiap kali parameter berubah, dari kuantil "
        "empiris 38 provinsi (ini perbandingan relatif, sengaja):",
    )
    formula(
        doc,
        "kelas = 0  jika nilai ≤ kuantil 1/3\n"
        "      = 1  jika nilai ≤ kuantil 2/3\n"
        "      = 2  jika di atas itu",
    )
    body(
        doc,
        "Tegak: tertil risiko (atas = tinggi). Datar: tertil pendidikan (kanan = tinggi).",
    )
    table(
        doc,
        ["Risiko", "Pendidikan", "Kuadran", "Arti singkat"],
        [
            ["tinggi", "rendah", "Senjang", "Prioritas intervensi pendidikan"],
            ["tinggi", "tinggi", "Selaras", "Simpul yang sudah sepadan"],
            ["rendah", "tinggi", "Berlebih", "Kapasitas lebih dari bebannya"],
            ["rendah", "rendah", "Relevan", "Sebanding pada beban kecil"],
            ["menengah", "menengah", "Menengah", "Bukan sudut matriks"],
        ],
    )

    heading(doc, "9. Indeks senjang")
    body(
        doc,
        "Footer peta mengurutkan enam provinsi dengan senjang terbesar. Untuk ranking "
        "ini, risiko dan IDPKI dinormalisasi min–maks 38 provinsi ke 0–1:",
    )
    formula(
        doc,
        "risiko_norm = (R − min R) / (max R − min R)\n"
        "IDPKI_norm  = (E − min E) / (max E − min E)\n\n"
        "senjang = risiko_norm − IDPKI_norm",
    )
    body(
        doc,
        "Mendekati +1: risiko relatif tinggi, pendidikan relatif rendah. Negatif: "
        "pendidikan relatif lebih tinggi daripada risiko.",
    )

    heading(doc, "Sumber data prototipe")
    bullets(
        doc,
        [
            "90 prodi S1 Teknik Sipil (BAN-PT / LAM Teknik).",
            "14 prodi kebencanaan (BAN-PT).",
            "Kurasi awal: 13 Planologi, 10 geologi, 8 arsitektur, 8 lingkungan, 8 kelautan, 9 multidisiplin.",
            "13 pusat studi, sekaligus proxy layanan kepakaran (daftar di bawah).",
            "Akreditasi institusi: prototipe. Internasional = ITB, UI, UGM, ITS; Unggul dikurasi; sisanya diinfer dari peringkat program.",
            "Penduduk: angka prototipe, 38 provinsi termasuk pemekaran Papua.",
            "Batas administrasi GeoJSON 38 provinsi.",
        ],
    )

    heading(doc, "Pusat studi dalam basis", level=2)
    table(
        doc,
        ["Pusat", "Perguruan tinggi", "Provinsi", "Fokus"],
        [[n, u, p, f] for n, u, p, f in CENTERS],
    )

    heading(doc, "Yang tidak diklaim")
    body(
        doc,
        "PKPRB bukan indeks kesiapsiagaan, bukan ranking kampus, dan bukan unduhan "
        "resmi IRBI per kabupaten. Bobot dapat digeser; hasil berubah. Inventaris "
        "prodi non-sipil dan akreditasi institusi perlu validasi BAN-PT. Skor bahaya "
        "selain komposit bersifat profil kerja, bukan angka BNPB.",
    )
    body(
        doc,
        "P2MI Multidisiplin FTSL ITB 2026 — Mainstreaming Disaster Resiliency in "
        "Infrastructure Systems.",
        italic=True,
        size=10,
    )

    out_public = Path("/workspace/public/Metodologi-PKPRB.docx")
    out_docs = Path("/workspace/docs/Metodologi-PKPRB.docx")
    out_docs.parent.mkdir(exist_ok=True)
    doc.save(out_public)
    doc.save(out_docs)
    print(out_public, out_public.stat().st_size)
    print(out_docs, out_docs.stat().st_size)


if __name__ == "__main__":
    build()
